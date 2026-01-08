"""
Serviço de Análise de Contratos com Inteligência Artificial
Suporta extração de texto de PDFs e documentos Word,
e análise com OpenAI GPT.
"""

import json
import logging
import os
import re
from decimal import Decimal
from datetime import datetime
from typing import Optional, Dict, Any

from django.conf import settings
from django.db.models import Q

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extrai texto de documentos PDF e Word"""
    
    @staticmethod
    def sanitize_text(text: str) -> str:
        """
        Remove caracteres problemáticos do texto extraído
        - Caracteres nulos (0x00)
        - Outros caracteres de controle problemáticos
        """
        if not text:
            return ""
        
        # Remove caracteres nulos e outros caracteres de controle problemáticos
        # Mantém apenas caracteres imprimíveis e quebras de linha válidas
        sanitized = text.replace('\x00', '')  # Remove caracteres nulos
        sanitized = sanitized.replace('\x01', '')  # Remove SOH
        sanitized = sanitized.replace('\x02', '')  # Remove STX
        sanitized = sanitized.replace('\x03', '')  # Remove ETX
        sanitized = sanitized.replace('\x04', '')  # Remove EOT
        
        # Remove outros caracteres de controle (exceto \n, \r, \t)
        # Mantém apenas caracteres imprimíveis, espaços, quebras de linha e tabs
        sanitized = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', sanitized)
        
        return sanitized
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extrai texto de arquivo PDF"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        # Sanitiza o texto antes de adicionar
                        page_text = DocumentExtractor.sanitize_text(page_text)
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            # Fallback para PyPDF2
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        # Sanitiza o texto antes de adicionar
                        text = DocumentExtractor.sanitize_text(text)
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
            except Exception as e:
                logger.error(f"Erro ao extrair PDF com PyPDF2: {e}")
                raise
        except Exception as e:
            logger.error(f"Erro ao extrair PDF: {e}")
            raise
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extrai texto de arquivo Word (.docx)"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Sanitiza o texto antes de adicionar
                    text = DocumentExtractor.sanitize_text(paragraph.text)
                    text_parts.append(text)
            
            # Também extrai de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        # Sanitiza o texto antes de adicionar
                        row_text = DocumentExtractor.sanitize_text(row_text)
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Erro ao extrair DOCX: {e}")
            raise
    
    @classmethod
    def extract_text(cls, file_path: str) -> str:
        """Extrai texto baseado na extensão do arquivo"""
        ext = file_path.split('.')[-1].lower()
        
        text = ""
        if ext == 'pdf':
            text = cls.extract_from_pdf(file_path)
        elif ext in ['docx', 'doc']:
            text = cls.extract_from_docx(file_path)
        else:
            raise ValueError(f"Formato de arquivo não suportado: {ext}")
        
        # Sanitiza o texto final antes de retornar (garantia extra)
        return cls.sanitize_text(text)


class ContractAIAnalyzer:
    """Analisa contratos usando IA (OpenAI)"""
    
    # Prompt base para análise de contratos
    SYSTEM_PROMPT = """Você é um especialista em análise de contratos públicos e privados no Brasil.
Sua tarefa é extrair informações estruturadas de documentos contratuais de forma COMPLETA e PRECISA.

Analise o documento fornecido e extraia as seguintes informações em formato JSON:

{
    "cliente": {
        "nome_razao_social": "Nome completo da empresa/órgão contratante (OBRIGATÓRIO)",
        "nome_fantasia": "Nome fantasia se disponível",
        "cnpj_cpf": "CNPJ ou CPF (apenas números, 11 ou 14 dígitos)",
        "tipo_cliente": "publico ou privado",
        "tipo_pessoa": "fisica ou juridica",
        "natureza_juridica": "Natureza jurídica da empresa (ex: Sociedade Anônima, Autarquia, etc)",
        "inscricao_estadual": "Inscrição Estadual (IE) se disponível",
        "inscricao_municipal": "Inscrição Municipal se disponível",
        "endereco": "Logradouro/rua (sem número)",
        "numero": "Número do endereço",
        "complemento": "Complemento do endereço (apto, sala, etc)",
        "bairro": "Bairro",
        "cidade": "Cidade",
        "estado": "UF (2 letras, ex: SP, RJ, DF)",
        "cep": "CEP (formato: 00000-000 ou apenas números)",
        "pais": "País (padrão: Brasil)",
        "nome_responsavel": "Nome do responsável/signatário legal",
        "cargo_responsavel": "Cargo do responsável",
        "email_contato": "Email principal de contato",
        "telefone_contato": "Telefone principal de contato"
    },
    "contrato": {
        "numero_contrato": "Número do contrato (OBRIGATÓRIO, ex: 001/2025, 123/2024)",
        "objeto": "Descrição completa e detalhada do objeto do contrato",
        "regime_legal": "LEI_14133 ou LEI_13303 ou PRIVADO",
        "modalidade_licitacao": "Modalidade de licitação (Pregão Eletrônico, Tomada de Preços, Concorrência, Dispensa, Inexigibilidade, etc)",
        "numero_processo": "Número do processo licitatório/administrativo",
        "numero_edital": "Número do edital de licitação se disponível",
        "pregao_eletronico": "Número do pregão eletrônico se aplicável",
        "ata_registro_preco": "Número da ARP (Ata de Registro de Preços) se aplicável",
        "orgao_gerenciador_arp": "Órgão gerenciador da ARP se aplicável",
        "termo_referencia": "Número do Termo de Referência (TR) se disponível",
        "numero_rfp_rfq": "Número do RFP/RFQ se for contrato privado",
        "data_assinatura": "YYYY-MM-DD (data de assinatura do contrato)",
        "vigencia_meses": "Número de meses de vigência (12, 24, 36, 48, 60 ou 120)",
        "valor_inicial": "Valor total do contrato em decimal (ex: 150000.00)",
        "fornecedores": ["Lista completa de fornecedores/fabricantes mencionados no contrato"],
        "origem_contrato": "LIC_14133_PROPRIA | ARP_GERENCIADOR | ARP_PARTICIPANTE | ARP_ADESAO_CARONA | DISPENSA_14133 | INEXIGIBILIDADE_14133 | LIC_13303_PROPRIA | CONTR_ESTATAL_DIRETA | RFP_PRIVADA | RFQ_PRIVADA | NEGOCIACAO_DIRETA_PRIVADA | FRAMEWORK_PRIVADO | OUTRO",
        "origem_contrato_justificativa": "Texto detalhado explicando por que esta origem foi escolhida com base nas cláusulas e dispositivos legais do contrato",
        "origem_contrato_confianca": 0.0
    },
    "itens": [
        {
            "lote": "Número do lote (padrão: 1 se não especificado)",
            "numero_item": "Número do item conforme documento (OBRIGATÓRIO)",
            "descricao": "Descrição completa e detalhada do item/serviço/produto",
            "tipo": "hardware, software, solucao, servico ou treinamento (IDENTIFICAR CORRETAMENTE baseado na descrição)",
            "unidade": "Unidade de medida (UN, HORA, MES, ANO, etc)",
            "quantidade": "Quantidade total (número decimal)",
            "valor_unitario": "Valor unitário em decimal (ex: 1500.00)",
            "vigencia_produto": "Vigência do produto em meses se aplicável (12, 24, 36, 48, 60, 120 ou null)"
        }
    ],
    "slas": [
        {
            "titulo": "Título do SLA (ex: Disponibilidade do Sistema, Tempo de Resposta, etc)",
            "descricao": "Descrição completa do SLA, incluindo condições e critérios",
            "tipo": "disponibilidade, tempo_resposta, resolucao, atendimento ou outro",
            "tempo_resposta_horas": "Tempo de resposta em horas (se aplicável)",
            "tempo_solucao_horas": "Tempo de solução/resolução em horas (se aplicável)",
            "percentual_disponibilidade": "Percentual de disponibilidade exigido (ex: 99.9)",
            "penalidade_percentual": "Percentual de penalidade/glosa se não cumprido",
            "observacoes": "Observações adicionais sobre o SLA"
        }
    ],
    "contatos_cliente": [
        {
            "nome": "Nome do contato",
            "email": "Email do contato",
            "telefone": "Telefone do contato",
            "funcao": "Função/cargo do contato"
        }
    ],
    "registros_existentes": {
        "cliente_existente": {
            "id": "ID do cliente se já existir no sistema (null se não existir)",
            "nome": "Nome do cliente existente",
            "cnpj_cpf": "CNPJ/CPF do cliente existente",
            "corresponde": true
        },
        "contrato_existente": {
            "id": "ID do contrato se já existir no sistema (null se não existir)",
            "numero_contrato": "Número do contrato existente",
            "corresponde": true
        },
        "itens_existentes": [
            {
                "id": "ID do item se já existir",
                "numero_item": "Número do item existente",
                "corresponde": true
            }
        ],
        "contatos_existentes": [
            {
                "id": "ID do contato se já existir",
                "nome": "Nome do contato existente",
                "email": "Email do contato existente",
                "corresponde": true
            }
        ]
    },
    "diario_bordo": {
        "resumo_geral": "Resumo executivo completo do contrato (2-3 parágrafos destacando objetivo principal, valor, vigência, principais entregas e características relevantes)",
        "itens_mais_importantes": [
            {
                "numero_item": "Número do item",
                "descricao": "Descrição resumida",
                "motivo_importancia": "Por que este item é importante para o contrato"
            }
        ],
        "operacionalizacao_equipe": "Texto detalhado explicando como o contrato deve ser operacionalizado pela equipe, incluindo processos, fluxos de trabalho, responsabilidades e procedimentos principais",
        "informacoes_gerente_contrato_cs": {
            "pontos_atencao": [
                {
                    "titulo": "Título do ponto de atenção",
                    "descricao": "Descrição detalhada",
                    "prioridade": "critica, alta, media ou baixa",
                    "acao_necessaria": "Ação específica a ser tomada"
                }
            ],
            "clausulas_criticas": [
                {
                    "numero_clausula": "Número da cláusula se disponível",
                    "titulo": "Título da cláusula",
                    "descricao": "Descrição completa da cláusula e seu impacto",
                    "prazo_atencao": "YYYY-MM-DD ou null se não houver prazo específico"
                }
            ],
            "slas_criticos": [
                {
                    "nome": "Nome do SLA crítico",
                    "descricao": "Descrição do SLA e por que é crítico",
                    "penalizacoes": "Descrição das penalizações/glosas se não cumprido"
                }
            ],
            "informacoes_comerciais": "Informações relevantes para gestão comercial e customer success (relacionamento com cliente, expectativas, pontos de atenção no relacionamento)"
        },
        "informacoes_gerente_projetos": {
            "escopo_execucao": "Descrição detalhada do escopo de execução dos serviços",
            "entregaveis_principais": [
                {
                    "nome": "Nome do entregável",
                    "descricao": "Descrição detalhada",
                    "prazo_estimado": "Prazo estimado ou referência temporal",
                    "dependencias": "Dependências de outros entregáveis ou atividades"
                }
            ],
            "requisitos_tecnicos": "Requisitos técnicos essenciais para execução dos serviços",
            "recursos_necessarios": "Recursos necessários (humanos, tecnológicos, materiais) para execução",
            "riscos_execucao": [
                {
                    "risco": "Descrição do risco",
                    "impacto": "Impacto potencial",
                    "mitigacao": "Estratégia de mitigação sugerida"
                }
            ],
            "cronograma_sugerido": "Sugestão de cronograma de execução baseado nas cláusulas e prazos do contrato"
        }
    },
    "observacoes": "Informações adicionais relevantes não capturadas nos campos acima",
    "confianca": {
        "geral": "alta, media ou baixa",
        "campos_incertos": ["Lista de campos com baixa confiança na extração"]
    }
}

REGRAS IMPORTANTES:
1. Se uma informação não estiver disponível, use null
2. Datas devem estar no formato YYYY-MM-DD
3. Valores monetários devem ser números decimais sem formatação
4. CNPJs devem conter apenas números (14 dígitos)
5. Para regime_legal, identifique pela menção às leis ou contexto
6. Para origem_contrato, escolha APENAS UM dos códigos permitidos e utilize o campo origem_contrato_justificativa para explicar sua decisão.
7. origem_contrato_confianca deve ser um número entre 0 e 1 indicando o nível de confiança na classificação da origem.
8. VERIFICAÇÃO DE REGISTROS EXISTENTES:
   - Se houver uma seção "REGISTROS EXISTENTES NO SISTEMA" no prompt, compare os dados extraídos com os registros listados.
   - Para cliente: compare CNPJ/CPF e nome. Se corresponder, preencha "registros_existentes.cliente_existente" com o ID e "corresponde": true.
   - Para contrato: compare número do contrato. Se corresponder, preencha "registros_existentes.contrato_existente" com o ID e "corresponde": true.
   - Para itens: compare número do item e descrição. Se corresponder, adicione em "registros_existentes.itens_existentes".
   - Para contatos: compare nome e email. Se corresponder, adicione em "registros_existentes.contatos_existentes".
   - Se não houver correspondência, deixe os campos como null ou array vazio.
9. CONTATOS DO CLIENTE:
   - Extraia TODOS os contatos mencionados no documento (não apenas o responsável).
   - Inclua nome, email, telefone e função/cargo de cada contato identificado.

10. IDENTIFICAÇÃO DO TIPO DE ITEM:
   - Analise cuidadosamente a descrição de cada item para identificar corretamente o tipo:
     * "hardware": Equipamentos físicos (servidores, switches, roteadores, computadores, etc)
     * "software": Licenças de software, aplicativos, sistemas
     * "solucao": Soluções completas que combinam hardware e software
     * "servico": Serviços de consultoria, suporte, manutenção, implementação
     * "treinamento": Cursos, capacitações, treinamentos
   - Seja preciso na classificação baseado na descrição real do item.

11. SLAs:
   - Extraia TODOS os SLAs mencionados no contrato.
   - Identifique corretamente o tipo (disponibilidade, tempo_resposta, resolucao, atendimento, outro).
   - Inclua valores de meta, penalizações e condições específicas.
   - Se houver percentuais de disponibilidade, inclua em "percentual_disponibilidade".

12. DIÁRIO DE BORDO:
   - O campo "diario_bordo" contém informações estratégicas para diferentes perfis:
     * "resumo_geral": Visão executiva do contrato
     * "itens_mais_importantes": Itens críticos que merecem atenção especial
     * "operacionalizacao_equipe": Como a equipe deve trabalhar com este contrato
     * "informacoes_gerente_contrato_cs": Informações para Gerente de Contrato e Customer Success
     * "informacoes_gerente_projetos": Informações técnicas para Gerente de Projetos
   - Seja detalhado e específico em cada seção, baseando-se nas cláusulas e termos do contrato.

13. Retorne APENAS o JSON, sem texto adicional
"""

    def __init__(self):
        """
        Inicializa o analisador de IA usando OpenAI
        """
        self._client = None
    
    def _get_openai_client(self):
        """Retorna cliente OpenAI configurado"""
        if self._client is None:
            try:
                from openai import OpenAI
                from decouple import config
                
                # Tenta ler do settings primeiro, depois do .env usando decouple
                api_key = getattr(settings, 'OPENAI_API_KEY', None)
                if not api_key:
                    api_key = config('OPENAI_API_KEY', default=None)
                if not api_key:
                    api_key = os.environ.get('OPENAI_API_KEY')
                
                if not api_key:
                    raise ValueError("OPENAI_API_KEY não configurada. Configure no arquivo .env ou nas variáveis de ambiente.")
                self._client = OpenAI(api_key=api_key)
            except ImportError:
                raise ImportError("Biblioteca openai não instalada. Execute: pip install openai")
        return self._client
    
    def analyze_with_openai(self, text: str, registros_existentes: Dict[str, Any] = None, model: str = "gpt-4o") -> Dict[str, Any]:
        """Analisa texto usando OpenAI GPT"""
        client = self._get_openai_client()
        
        # Limita o texto para não exceder contexto
        max_chars = 100000  # ~25k tokens
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[... documento truncado ...]"
        
        # Monta prompt com informações de registros existentes
        user_prompt = f"Analise o seguinte documento:\n\n{text}"
        
        if registros_existentes:
            registros_info = self._formatar_registros_existentes(registros_existentes)
            user_prompt += f"\n\n=== REGISTROS EXISTENTES NO SISTEMA ===\n{registros_info}\n\nIMPORTANTE: Verifique se os dados extraídos do documento correspondem a algum dos registros acima. Se corresponder, indique no campo 'registro_existente' do JSON."
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": self.SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content
        return json.loads(result)
    
    def _formatar_registros_existentes(self, registros: Dict[str, Any]) -> str:
        """Formata informações de registros existentes para incluir no prompt"""
        texto = ""
        
        if registros.get('clientes'):
            texto += "CLIENTES CADASTRADOS:\n"
            for cliente in registros['clientes']:
                texto += f"- ID: {cliente['id']}, Nome: {cliente['nome_razao_social']}, CNPJ/CPF: {cliente['cnpj_cpf']}, Tipo: {cliente['tipo_cliente']}\n"
                if cliente.get('contatos'):
                    texto += "  Contatos:\n"
                    for contato in cliente['contatos']:
                        texto += f"    - {contato['nome']} ({contato['email']}, {contato['telefone']}) - {contato.get('funcao', 'N/A')}\n"
        
        if registros.get('contratos'):
            texto += "\nCONTRATOS CADASTRADOS:\n"
            for contrato in registros['contratos']:
                texto += f"- ID: {contrato['id']}, Número: {contrato['numero_contrato']}, Cliente: {contrato['cliente']}\n"
                texto += f"  Objeto: {contrato['objeto']}\n"
                if contrato.get('itens'):
                    texto += "  Itens:\n"
                    for item in contrato['itens']:
                        texto += f"    - Item {item['numero_item']}: {item['descricao']} (Tipo: {item['tipo']})\n"
        
        return texto
    
    def analyze(self, text: str, registros_existentes: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Analisa texto usando OpenAI GPT
        
        Args:
            text: Texto a ser analisado
            registros_existentes: Dict com informações de registros existentes no sistema
        """
        return self.analyze_with_openai(text, registros_existentes)


class ContractAIService:
    """
    Serviço principal para processamento de documentos de contrato
    """
    
    def __init__(self):
        self.extractor = DocumentExtractor()
        self.analyzer = ContractAIAnalyzer()
    
    def process_multiple_documents(self, analise) -> Dict[str, Any]:
        """
        Processa múltiplos documentos de uma análise
        
        Args:
            analise: Instância de AnaliseContrato
            
        Returns:
            Dict com dados extraídos consolidados
        """
        from contracts.models import AnaliseContrato, Cliente, ContatoCliente, Contrato, ItemContrato
        
        try:
            # Atualiza status
            analise.status = 'processando'
            analise.save(update_fields=['status'])
            
            # Extrai texto de todos os documentos
            textos_por_tipo = {}
            textos_consolidados = []
            
            for documento in analise.documentos.all():
                try:
                    documento.status = 'processando'
                    documento.save(update_fields=['status'])
                    
                    file_path = documento.arquivo.path
                    texto = self.extractor.extract_text(file_path)
                    # Sanitiza novamente antes de salvar (garantia extra)
                    texto = DocumentExtractor.sanitize_text(texto)
                    documento.texto_extraido = texto
                    documento.status = 'analisado'
                    documento.save(update_fields=['texto_extraido', 'status'])
                    
                    # Organiza por tipo de documento
                    tipo = documento.tipo_documento
                    if tipo not in textos_por_tipo:
                        textos_por_tipo[tipo] = []
                    textos_por_tipo[tipo].append({
                        'nome': documento.nome,
                        'texto': texto
                    })
                    
                    textos_consolidados.append(f"=== {documento.get_tipo_documento_display()}: {documento.nome} ===\n{texto}")
                    
                except Exception as e:
                    logger.error(f"Erro ao processar documento {documento.id}: {e}")
                    documento.status = 'erro'
                    documento.mensagem_erro = str(e)
                    documento.save(update_fields=['status', 'mensagem_erro'])
            
            # Consolida todos os textos
            texto_consolidado = "\n\n".join(textos_consolidados)
            # Sanitiza o texto consolidado antes de salvar e analisar
            texto_consolidado = DocumentExtractor.sanitize_text(texto_consolidado)
            analise.texto_consolidado = texto_consolidado
            analise.save(update_fields=['texto_consolidado'])
            
            # Busca registros existentes para incluir no contexto da análise
            registros_existentes = self._buscar_registros_existentes(analise)
            
            # Analisa com IA usando todos os textos e informações de registros existentes
            dados = self.analyzer.analyze(texto_consolidado, registros_existentes)
            
            # Adiciona informações sobre os documentos processados
            dados['documentos_processados'] = {
                tipo: [doc['nome'] for doc in docs]
                for tipo, docs in textos_por_tipo.items()
            }
            
            # Adiciona informações sobre registros existentes encontrados
            dados['registros_existentes'] = registros_existentes
            
            analise.dados_extraidos = dados
            analise.status = 'analisado'
            analise.save(update_fields=['dados_extraidos', 'status'])
            
            return dados
            
        except Exception as e:
            logger.error(f"Erro ao processar análise {analise.id}: {e}")
            analise.status = 'erro'
            analise.mensagem_erro = str(e)
            analise.save(update_fields=['status', 'mensagem_erro'])
            raise
    
    def _buscar_registros_existentes(self, analise) -> Dict[str, Any]:
        """
        Busca registros existentes no sistema que podem estar relacionados à análise
        
        Args:
            analise: Instância de AnaliseContrato
            
        Returns:
            Dict com informações de registros existentes
        """
        from contracts.models import Cliente, ContatoCliente, Contrato, ItemContrato
        
        registros = {
            'clientes': [],
            'contratos': [],
            'itens_contrato': []
        }
        
        # Se já tem cliente/contrato vinculado, busca informações completas
        if analise.cliente_gerado:
            cliente = analise.cliente_gerado
            registros['clientes'].append({
                'id': cliente.id,
                'nome_razao_social': cliente.nome_razao_social,
                'cnpj_cpf': cliente.cnpj_cpf,
                'tipo_cliente': cliente.tipo_cliente,
                'contatos': [
                    {
                        'nome': contato.nome,
                        'email': contato.email,
                        'telefone': contato.telefone,
                        'funcao': contato.funcao
                    }
                    for contato in cliente.contatos.filter(ativo=True)
                ]
            })
        
        if analise.contrato_gerado:
            contrato = analise.contrato_gerado
            registros['contratos'].append({
                'id': contrato.id,
                'numero_contrato': contrato.numero_contrato,
                'cliente': contrato.cliente.nome_razao_social,
                'objeto': contrato.objeto[:200] if contrato.objeto else '',
                'itens': [
                    {
                        'numero_item': item.numero_item,
                        'descricao': item.descricao[:100],
                        'tipo': item.tipo
                    }
                    for item in contrato.itens.all()[:10]  # Limita a 10 itens
                ]
            })
        
        # Busca clientes similares por CNPJ parcial (se houver número de contrato na análise)
        # Isso ajuda a identificar possíveis matches mesmo sem CNPJ completo
        
        return registros
    
    @staticmethod
    def verificar_contrato_existente(numero_contrato: str = None, cnpj_cliente: str = None):
        """
        Verifica se um contrato já existe
        
        Args:
            numero_contrato: Número do contrato para buscar
            cnpj_cliente: CNPJ do cliente para buscar
            
        Returns:
            Tupla (contrato, cliente) ou (None, None) se não existir
        """
        from contracts.models import Contrato, Cliente
        
        contrato = None
        cliente = None
        
        # Busca por número do contrato
        if numero_contrato:
            contrato = Contrato.objects.filter(numero_contrato=numero_contrato).first()
            if contrato:
                cliente = contrato.cliente
        
        # Se não encontrou, busca por CNPJ do cliente
        if not contrato and cnpj_cliente:
            cnpj_limpo = cnpj_cliente.replace('.', '').replace('/', '').replace('-', '')
            cliente = Cliente.objects.filter(cnpj_cpf=cnpj_limpo).first()
            if cliente:
                # Busca o contrato mais recente do cliente
                contrato = cliente.contratos.order_by('-data_assinatura').first()
        
        return contrato, cliente
    
    @staticmethod
    def create_cliente_from_data(dados: Dict[str, Any], user=None, cliente_existente=None):
        """Cria ou retorna um Cliente a partir dos dados extraídos"""
        from contracts.models import Cliente
        
        # Se já existe um cliente, retorna ele
        if cliente_existente:
            return cliente_existente
        
        cliente_data = dados.get('cliente', {})
        if not cliente_data or not cliente_data.get('nome_razao_social'):
            return None
        
        # Verifica se já existe pelo CNPJ
        cnpj = cliente_data.get('cnpj_cpf', '').replace('.', '').replace('/', '').replace('-', '')
        if cnpj:
            existing = Cliente.objects.filter(cnpj_cpf=cnpj).first()
            if existing:
                return existing
        
        cliente = Cliente(
            nome_razao_social=cliente_data.get('nome_razao_social', ''),
            nome_fantasia=cliente_data.get('nome_fantasia') or None,
            tipo_cliente=cliente_data.get('tipo_cliente', 'publico'),
            tipo_pessoa=cliente_data.get('tipo_pessoa', 'juridica'),
            cnpj_cpf=cnpj or 'A_DEFINIR',
            natureza_juridica=cliente_data.get('natureza_juridica') or None,
            inscricao_estadual=cliente_data.get('inscricao_estadual') or None,
            inscricao_municipal=cliente_data.get('inscricao_municipal') or None,
            endereco=cliente_data.get('endereco') or 'A definir',
            numero=cliente_data.get('numero') or 'S/N',
            complemento=cliente_data.get('complemento') or None,
            bairro=cliente_data.get('bairro') or 'Centro',
            cidade=cliente_data.get('cidade') or 'A definir',
            estado=(cliente_data.get('estado') or 'DF')[:2],
            cep=cliente_data.get('cep') or '00000-000',
            pais=cliente_data.get('pais') or 'Brasil',
            nome_responsavel=cliente_data.get('nome_responsavel') or 'A definir',
            cargo_responsavel=cliente_data.get('cargo_responsavel') or 'A definir',
            telefone_contato=cliente_data.get('telefone_contato') or '(00) 0000-0000',
            email_contato=cliente_data.get('email_contato') or 'a@definir.com',
        )
        cliente.save()
        return cliente
    
    @staticmethod
    def create_contrato_from_data(dados: Dict[str, Any], cliente, user=None, contrato_existente=None):
        """Cria ou retorna um Contrato a partir dos dados extraídos"""
        from contracts.models import Contrato
        
        # Se já existe um contrato, retorna ele
        if contrato_existente:
            return contrato_existente
        
        contrato_data = dados.get('contrato', {})
        if not contrato_data:
            return None
        
        # Verifica se já existe pelo número do contrato
        numero_contrato = contrato_data.get('numero_contrato', '')
        if numero_contrato and numero_contrato != 'A_DEFINIR':
            existing = Contrato.objects.filter(numero_contrato=numero_contrato).first()
            if existing:
                return existing
        
        # Parse da data
        data_assinatura = None
        if contrato_data.get('data_assinatura'):
            try:
                data_assinatura = datetime.strptime(
                    contrato_data['data_assinatura'], '%Y-%m-%d'
                ).date()
            except:
                data_assinatura = datetime.now().date()
        
        # Valor inicial
        valor_inicial = Decimal('0.00')
        if contrato_data.get('valor_inicial'):
            try:
                valor_inicial = Decimal(str(contrato_data['valor_inicial']))
            except:
                pass
        
        # Regime legal
        regime_map = {
            'LEI_14133': 'LEI_14133',
            'LEI_13303': 'LEI_13303',
            'PRIVADO': 'PRIVADO',
        }
        regime = regime_map.get(contrato_data.get('regime_legal', '').upper(), 'LEI_14133')
        
        # Vigência - garantir que é um valor válido
        vigencia_meses = int(contrato_data.get('vigencia_meses', 12) or 12)
        vigencias_validas = [12, 24, 36, 48, 60, 120]
        if vigencia_meses not in vigencias_validas:
            # Arredonda para a vigência válida mais próxima
            vigencia_meses = min(vigencias_validas, key=lambda x: abs(x - vigencia_meses))
        
        # Origem do contrato (default: RFP_PRIVADA para contratos privados,
        # LIC_14133_PROPRIA para públicos quando não informado)
        origem_contrato = contrato_data.get('origem_contrato')
        origem_justificativa = contrato_data.get('origem_contrato_justificativa')
        origem_confianca = contrato_data.get('origem_contrato_confianca')

        from contracts.models import Contrato as ContratoModel

        if not origem_contrato:
            if regime == 'PRIVADO':
                origem_contrato = ContratoModel.OrigemContrato.RFP_PRIVADA
            elif regime == 'LEI_14133':
                origem_contrato = ContratoModel.OrigemContrato.LIC_14133_PROPRIA
            elif regime == 'LEI_13303':
                origem_contrato = ContratoModel.OrigemContrato.LIC_13303_PROPRIA
            else:
                origem_contrato = ContratoModel.OrigemContrato.OUTRO

        contrato = Contrato(
            cliente=cliente,
            numero_contrato=contrato_data.get('numero_contrato', 'A_DEFINIR'),
            objeto=contrato_data.get('objeto', ''),
            regime_legal=regime,
            modalidade_licitacao=contrato_data.get('modalidade_licitacao') or None,
            numero_edital=contrato_data.get('numero_edital') or None,
            pregao_eletronico=contrato_data.get('pregao_eletronico') or None,
            processo=contrato_data.get('numero_processo') or None,
            termo_referencia=contrato_data.get('termo_referencia') or None,
            ata_registro_preco=contrato_data.get('ata_registro_preco') or None,
            orgao_gerenciador_arp=contrato_data.get('orgao_gerenciador_arp') or None,
            numero_rfp_rfq=contrato_data.get('numero_rfp_rfq') or None,
            data_assinatura=data_assinatura or datetime.now().date(),
            vigencia=vigencia_meses,
            valor_inicial=valor_inicial,
            fornecedores=contrato_data.get('fornecedores', []),
            origem_contrato=origem_contrato,
            origem_contrato_justificativa_ia=origem_justificativa,
            origem_contrato_confianca_ia=origem_confianca if origem_confianca is not None else None,
        )
        contrato.save()
        return contrato
    
    @staticmethod
    def create_itens_from_data(dados: Dict[str, Any], contrato):
        """Cria Itens de Contrato a partir dos dados extraídos"""
        from contracts.models import ItemContrato
        
        itens_data = dados.get('itens', [])
        itens_criados = []
        
        tipo_map = {
            'hardware': 'hardware',
            'software': 'software',
            'solucao': 'solucao',
            'servico': 'servico',
            'treinamento': 'treinamento',
        }
        
        for item_data in itens_data:
            try:
                quantidade = Decimal(str(item_data.get('quantidade', 1) or 1))
                valor_unitario = Decimal(str(item_data.get('valor_unitario', 0) or 0))
                
                vigencia_produto = item_data.get('vigencia_produto')
                if vigencia_produto:
                    try:
                        vigencia_produto = int(vigencia_produto)
                        vigencias_validas = [12, 24, 36, 48, 60, 120]
                        if vigencia_produto not in vigencias_validas:
                            vigencia_produto = min(vigencias_validas, key=lambda x: abs(x - vigencia_produto))
                    except:
                        vigencia_produto = None
                else:
                    vigencia_produto = None
                
                item = ItemContrato(
                    contrato=contrato,
                    lote=int(item_data.get('lote', 1) or 1),
                    numero_item=str(item_data.get('numero_item', len(itens_criados) + 1)),
                    descricao=item_data.get('descricao', 'Item sem descrição'),
                    tipo=tipo_map.get(item_data.get('tipo', '').lower(), 'servico'),
                    unidade=item_data.get('unidade', 'UN'),
                    quantidade=quantidade,
                    valor_unitario=valor_unitario,
                    vigencia_produto=vigencia_produto,
                )
                item.save()
                itens_criados.append(item)
            except Exception as e:
                logger.warning(f"Erro ao criar item: {e}")
                continue
        
        return itens_criados
    
    @staticmethod
    def create_slas_from_data(dados: Dict[str, Any], contrato):
        """Cria SLAs a partir dos dados extraídos"""
        from contracts.models import SLA
        
        slas_data = dados.get('slas', [])
        slas_criados = []
        
        tipo_map = {
            'disponibilidade': 'disponibilidade',
            'tempo_resposta': 'tempo_resposta',
            'resolucao': 'resolucao',
            'atendimento': 'atendimento',
            'outro': 'outro',
        }
        
        for sla_data in slas_data:
            try:
                titulo = sla_data.get('titulo') or sla_data.get('nome', 'SLA')
                tipo = tipo_map.get(sla_data.get('tipo', '').lower(), 'outro')
                
                # Monta a meta do SLA
                meta_parts = []
                if sla_data.get('tempo_resposta_horas'):
                    meta_parts.append(f"{sla_data.get('tempo_resposta_horas')} horas de resposta")
                if sla_data.get('tempo_solucao_horas'):
                    meta_parts.append(f"{sla_data.get('tempo_solucao_horas')} horas de solução")
                if sla_data.get('percentual_disponibilidade'):
                    meta_parts.append(f"{sla_data.get('percentual_disponibilidade')}% de disponibilidade")
                meta = ', '.join(meta_parts) if meta_parts else 'A definir'
                
                # Valor da penalidade
                penalidade_percentual = sla_data.get('penalidade_percentual')
                valor_penalidade = None
                if penalidade_percentual:
                    try:
                        # Calcula valor da penalidade como percentual do valor do contrato
                        if contrato.valor_inicial:
                            valor_penalidade = contrato.valor_inicial * (Decimal(str(penalidade_percentual)) / 100)
                    except:
                        pass
                
                sla = SLA(
                    contrato=contrato,
                    titulo=titulo,
                    descricao=sla_data.get('descricao', ''),
                    tipo=tipo,
                    meta=meta,
                    valor_penalidade=valor_penalidade,
                )
                sla.save()
                slas_criados.append(sla)
            except Exception as e:
                logger.warning(f"Erro ao criar SLA: {e}")
                continue
        
        return slas_criados
    
    @staticmethod
    def create_diario_bordo_from_data(dados: Dict[str, Any], contrato, user=None):
        """Cria o primeiro registro do Diário de Bordo a partir dos dados extraídos pela IA"""
        from contracts.models import DiarioBordoContrato
        
        diario_data = dados.get('diario_bordo', {})
        if not diario_data:
            return None
        
        # Monta o conteúdo completo do registro
        conteudo_parts = []
        
        if diario_data.get('resumo_geral'):
            conteudo_parts.append(f"## Resumo Geral\n\n{diario_data.get('resumo_geral')}\n")
        
        if diario_data.get('operacionalizacao_equipe'):
            conteudo_parts.append(f"## Operacionalização pela Equipe\n\n{diario_data.get('operacionalizacao_equipe')}\n")
        
        conteudo = '\n'.join(conteudo_parts) if conteudo_parts else 'Análise inicial do contrato realizada pela IA.'
        
        diario = DiarioBordoContrato(
            contrato=contrato,
            tipo_registro='analise_ia',
            titulo='Análise Inicial do Contrato - IA',
            conteudo=conteudo,
            resumo_geral=diario_data.get('resumo_geral') or None,
            itens_mais_importantes=diario_data.get('itens_mais_importantes', []),
            operacionalizacao_equipe=diario_data.get('operacionalizacao_equipe') or None,
            informacoes_gerente_contrato_cs=diario_data.get('informacoes_gerente_contrato_cs', {}),
            informacoes_gerente_projetos=diario_data.get('informacoes_gerente_projetos', {}),
            criado_por=user,
        )
        diario.save()
        return diario
    
    # Prompt otimizado com melhores práticas de gestão de projetos (PMBOK, Scrum, Agile)
    PLANO_TRABALHO_PROMPT = """Você é especialista em gestão de projetos de infraestrutura de TI e cibersegurança, certificado (PMP/PMI, Scrum Master, ITIL, CISSP) com expertise em metodologias ágeis aplicadas a projetos de infraestrutura e segurança da informação.

ESTRUTURA DO PROJETO:
- PROJETO: Contém múltiplas SPRINTS e TAREFAS
- SPRINT: Ciclo de desenvolvimento que agrupa TAREFAS relacionadas
- TAREFA: Unidade de trabalho que pode estar em uma SPRINT ou no BACKLOG do projeto (sem sprint atribuída)
- BACKLOG DO PROJETO: Tarefas sem sprint atribuída que aguardam alocação

IMPORTANTE: 
- Tarefas podem ser criadas sem sprint (ficam no backlog do projeto)
- Tarefas precisam ser atribuídas a uma sprint para serem iniciadas/executadas
- Tarefas precisam ter um responsável para serem executadas
- Tarefas podem ter uma flag para bilhetar ou não na Ordem de Serviço (OS) do projeto

Analise o contrato e documentos fornecidos, com FOCO ESPECIAL nos ITENS DO CONTRATO (Serviço ou Treinamento) vinculados ao projeto, aplicando as melhores práticas de gestão ágil de projetos de infraestrutura de TI e cibersegurança.

MELHORES PRÁTICAS A APLICAR (Infraestrutura de TI e Cibersegurança):
1. ITIL 4: Framework de gerenciamento de serviços de TI (Service Strategy, Design, Transition, Operation, Continual Improvement)
2. PMBOK 7ª edição: 5 grupos de processos adaptados para projetos de infraestrutura
3. Scrum/Agile: Metodologias ágeis adaptadas para projetos de infraestrutura (DevOps, SRE)
4. Cibersegurança: Frameworks NIST, ISO 27001, CIS Controls
5. Gestão de Infraestrutura: Planejamento de capacidade, disponibilidade, continuidade de negócios
6. Gestão de Riscos de Segurança: Identificação de vulnerabilidades, análise de ameaças, planos de mitigação
7. Gestão de Mudanças em Produção: Processo controlado para mudanças em ambiente crítico
8. Gestão de Stakeholders: Identificação, engajamento e comunicação com equipes técnicas e de negócio

Gere JSON estruturado seguindo este formato:

{
    "resumo_contrato": "Resumo executivo do contrato (2-3 parágrafos destacando objetivo, escopo, valor e principais entregas)",
    "clausulas_criticas": [
        {
            "titulo": "Título da cláusula crítica",
            "numero_clausula": "Número da cláusula se disponível",
            "descricao": "Descrição detalhada da cláusula e seu impacto no projeto",
            "impacto": "critico, alto ou medio",
            "acao_necessaria": "Ações específicas para cumprir ou mitigar riscos",
            "prazo_atencao": "YYYY-MM-DD ou null",
            "risco_associado": "Descrição do risco se não cumprida"
        }
    ],
    "pontos_atencao": [
        {
            "titulo": "Título do ponto de atenção",
            "descricao": "Descrição detalhada do ponto crítico",
            "prioridade": "critica, alta, media ou baixa",
            "acao_recomendada": "Ação específica e acionável para mitigar",
            "responsavel": "Quem deve executar a ação",
            "prazo": "YYYY-MM-DD ou null"
        }
    ],
    "slas_importantes": [
        {
            "nome": "Nome do SLA (específico e mensurável)",
            "descricao": "Descrição detalhada do SLA e seu contexto",
            "tempo_resposta_horas": 4,
            "tempo_solucao_horas": 24,
            "prioridade": "critica, alta, media ou baixa",
            "alerta_antes_horas": 24,
            "penalizacoes": [
                {
                    "descricao": "Descrição clara da penalização/glosa",
                    "tipo": "penalizacao, glosa ou multa",
                    "percentual": 5.0,
                    "valor_fixo": null,
                    "condicao_aplicacao": "Condições específicas de aplicação"
                }
            ],
            "indicadores": [
                {
                    "nome": "Nome do indicador (KPI)",
                    "meta": "Meta quantificável a ser atingida",
                    "unidade": "Unidade de medida (%, horas, dias, etc)",
                    "frequencia_medicao": "diaria, semanal, mensal"
                }
            ]
        }
    ],
    "matriz_raci": [
        {
            "atividade": "Nome da atividade/entregável",
            "fase": "planejamento, implantacao, execucao ou suporte",
            "responsavel": "Quem executa (R) - nome específico ou função",
            "aprovador": "Quem aprova (A) - nome específico ou função",
            "consultado": "Quem é consultado (C), separado por vírgula",
            "informado": "Quem é informado (I), separado por vírgula"
        }
    ],
    "fluxo_trabalho_fases": [
        {
            "fase": "Planejamento",
            "descricao": "Fase de iniciação e planejamento detalhado do projeto (PMBOK: Iniciação + Planejamento)",
            "duracao_dias": 30,
            "objetivo_fase": "Objetivo claro e mensurável desta fase",
            "entregaveis_principais": ["Lista de entregáveis principais da fase"],
            "sprints": [
                {
                    "nome": "Sprint 1: Kick-off e Definição de Escopo",
                    "objetivo": "Objetivo específico e mensurável da sprint (Sprint Goal)",
                    "data_inicio": "YYYY-MM-DD",
                    "data_fim": "YYYY-MM-DD",
                    "duracao_semanas": 2,
                    "entregaveis": ["Lista de entregáveis da sprint"],
                    "tarefas": [
                        {
                            "titulo": "Título específico e acionável da tarefa",
                            "descricao": "Descrição detalhada incluindo critérios de aceitação",
                            "tipo": "planejamento, desenvolvimento, teste, documentacao, treinamento, implantacao, execucao ou suporte",
                            "prioridade": "critica, alta, media ou baixa",
                            "horas_planejadas": 40,
                            "dependencias": ["Título de tarefas que devem ser concluídas antes"],
                            "entregavel": "Entregável específico desta tarefa",
                            "criterios_aceitacao": "Critérios claros de aceitação",
                            "responsavel": "Função ou papel responsável"
                        }
                    ]
                }
            ]
        },
        {
            "fase": "Implantação",
            "descricao": "Fase de implantação da solução (PMBOK: Execução com foco em implementação)",
            "duracao_dias": 60,
            "objetivo_fase": "Objetivo claro e mensurável desta fase",
            "entregaveis_principais": ["Lista de entregáveis principais da fase"],
            "sprints": [
                {
                    "nome": "Sprint 1: Configuração de Ambiente e Infraestrutura",
                    "objetivo": "Objetivo específico e mensurável da sprint",
                    "data_inicio": "YYYY-MM-DD",
                    "data_fim": "YYYY-MM-DD",
                    "duracao_semanas": 2,
                    "entregaveis": ["Lista de entregáveis da sprint"],
                    "tarefas": [
                        {
                            "titulo": "Configurar ambiente de desenvolvimento",
                            "descricao": "Descrição detalhada com especificações técnicas",
                            "tipo": "implantacao",
                            "prioridade": "alta",
                            "horas_planejadas": 40,
                            "dependencias": [],
                            "entregavel": "Ambiente configurado e documentado",
                            "criterios_aceitacao": "Ambiente funcional, documentação completa, testes de conectividade OK",
                            "responsavel": "Equipe de Infraestrutura"
                        }
                    ]
                }
            ]
        },
        {
            "fase": "Execução",
            "descricao": "Fase de execução e operação (PMBOK: Execução + Monitoramento/Controle)",
            "duracao_dias": 90,
            "objetivo_fase": "Objetivo claro e mensurável desta fase",
            "entregaveis_principais": ["Lista de entregáveis principais da fase"],
            "sprints": [
                {
                    "nome": "Sprint 1: Operação Assistida e Entrega Contínua",
                    "objetivo": "Objetivo específico e mensurável da sprint",
                    "data_inicio": "YYYY-MM-DD",
                    "data_fim": "YYYY-MM-DD",
                    "duracao_semanas": 2,
                    "entregaveis": ["Lista de entregáveis da sprint"],
                    "tarefas": [
                        {
                            "titulo": "Monitoramento e suporte operacional",
                            "descricao": "Descrição detalhada das atividades",
                            "tipo": "execucao",
                            "prioridade": "alta",
                            "horas_planejadas": 40,
                            "dependencias": [],
                            "entregavel": "Relatório de operação e métricas",
                            "criterios_aceitacao": "Sistema operacional, SLAs atendidos, relatórios gerados",
                            "responsavel": "Equipe de Operações"
                        }
                    ]
                }
            ]
        },
        {
            "fase": "Suporte/Sustentação",
            "descricao": "Fase de suporte, manutenção e encerramento (PMBOK: Encerramento + Operações)",
            "duracao_dias": 180,
            "objetivo_fase": "Objetivo claro e mensurável desta fase",
            "entregaveis_principais": ["Lista de entregáveis principais da fase"],
            "sprints": [
                {
                    "nome": "Sprint 1: Sustentação e Melhorias Contínuas",
                    "objetivo": "Objetivo específico e mensurável da sprint",
                    "data_inicio": "YYYY-MM-DD",
                    "data_fim": "YYYY-MM-DD",
                    "duracao_semanas": 2,
                    "entregaveis": ["Lista de entregáveis da sprint"],
                    "tarefas": [
                        {
                            "titulo": "Suporte técnico e atendimento a chamados",
                            "descricao": "Descrição detalhada das atividades de suporte",
                            "tipo": "suporte",
                            "prioridade": "media",
                            "horas_planejadas": 40,
                            "dependencias": [],
                            "entregavel": "Chamados resolvidos conforme SLA",
                            "criterios_aceitacao": "Tempo de resposta e solução dentro dos SLAs, satisfação do cliente",
                            "responsavel": "Equipe de Suporte"
                        }
                    ]
                }
            ]
        }
    ],
    "cronograma": {
        "data_inicio_prevista": "YYYY-MM-DD",
        "data_fim_prevista": "YYYY-MM-DD",
        "marcos": [
            {
                "nome": "Nome do marco (Milestone)",
                "data": "YYYY-MM-DD",
                "descricao": "Descrição do marco e entregáveis associados",
                "tipo": "inicio, planejamento, implantacao, execucao, suporte ou encerramento",
                "entregaveis": ["Lista de entregáveis do marco"]
            }
        ],
        "buffer_risco_dias": 10
    },
    "plano_comunicacao": {
        "stakeholders": [
            {
                "nome": "Nome ou função do stakeholder",
                "papel": "Papel no projeto (Patrocinador, Gerente, Cliente, Equipe, etc)",
                "nivel_interesse": "alto, medio ou baixo",
                "nivel_influencia": "alto, medio ou baixo",
                "frequencia_comunicacao": "diaria, semanal, quinzenal ou mensal",
                "canais": ["email", "reuniao", "relatorio", "dashboard"],
                "tipo_informacao": "estratégica, tática ou operacional"
            }
        ],
        "reunioes": [
            {
                "tipo": "Tipo de reunião (Daily, Sprint Planning, Review, Retrospective, Steering Committee, etc)",
                "frequencia": "diaria, semanal, quinzenal ou mensal",
                "duracao_minutos": 60,
                "participantes": ["Lista de participantes ou funções"],
                "objetivo": "Objetivo específico da reunião",
                "agenda": "Itens típicos da agenda"
            }
        ]
    },
    "template_status_report": "Template estruturado para relatórios de status incluindo: Resumo Executivo, Progresso, Entregáveis, Riscos, Próximos Passos",
    "frequencia_status_report": "semanal ou mensal"
}

REGRAS E DIRETRIZES OBRIGATÓRIAS:

1. ESTRUTURA DO PLANO DE TRABALHO:
   - O plano é para um PROJETO específico vinculado a um ITEM DO CONTRATO (tipo Serviço ou Treinamento)
   - Estrutura: PROJETO → SPRINTS → TAREFAS
   - Tarefas podem ser criadas sem sprint (ficam no backlog do projeto)
   - Tarefas precisam ser atribuídas a uma sprint para serem iniciadas/executadas
   - Tarefas precisam ter um responsável para serem executadas
   - Tarefas podem ter flag para bilhetar ou não na Ordem de Serviço (OS) do projeto

2. FOCO NO ITEM DO CONTRATO VINCULADO AO PROJETO:
   - O plano deve ser estruturado considerando o ITEM DO CONTRATO específico vinculado ao projeto (tipo Serviço ou Treinamento)
   - Analise o nome do projeto, descrição e características do item do contrato para refinar o plano
   - O plano deve refletir as necessidades específicas do item (Serviço ou Treinamento)
   - Considere o modelo de plano de trabalho do fornecedor vinculado ao contrato

3. PROCESSO DE EXECUÇÃO (Sprints e Tarefas):
   - O plano deve criar SPRINTS que agrupam TAREFAS relacionadas
   - Cada sprint deve ter: nome, objetivo, data de início, data de fim, e lista de tarefas
   - Cada tarefa deve ter: título, descrição, horas planejadas, prioridade, responsável sugerido
   - Tarefas podem ser criadas sem sprint (ficam no backlog do projeto para alocação futura)
   - Tarefas atribuídas a sprints podem ser iniciadas/executadas quando tiverem responsável
   - Considere criar algumas tarefas no backlog (sem sprint) para alocação posterior

3. CRONOGRAMA:
   - Deve respeitar rigorosamente a vigência do contrato
   - Data de início: data de assinatura do contrato
   - Data de fim: data de término do contrato
   - Considerar dependências entre itens ao calcular durações
   - Incluir buffer de risco (10-15% do tempo total)

4. FASES POR ITEM (Infraestrutura de TI e Cibersegurança):
   - Planejamento: Análise de requisitos, arquitetura, planejamento de segurança, definição de escopo (10-20% do tempo)
   - Implantação: Configuração, instalação, hardening de segurança, testes de segurança (20-30% do tempo)
   - Execução: Operação, monitoramento de segurança, gestão de vulnerabilidades, compliance (30-40% do tempo)
   - Suporte: Suporte contínuo, atualizações de segurança, manutenção preventiva, gestão de incidentes (30-40% do tempo)

5. MELHORES PRÁTICAS DE INFRAESTRUTURA DE TI:
   - Considerar alta disponibilidade, redundância, backup e disaster recovery
   - Planejar capacidade e escalabilidade
   - Definir procedimentos de operação e manutenção
   - Incluir documentação técnica e operacional

6. MELHORES PRÁTICAS DE CIBERSEGURANÇA:
   - Identificar requisitos de segurança para cada item
   - Planejar hardening, configuração segura, gestão de acessos
   - Considerar compliance (LGPD, ISO 27001, etc.)
   - Incluir monitoramento de segurança, gestão de vulnerabilidades, resposta a incidentes

7. ENTREGÁVEIS:
   - Cada etapa deve ter entregáveis claros e mensuráveis
   - Considerar documentos técnicos, configurações, sistemas operacionais, treinamentos
   - Entregáveis devem ser verificáveis e aceitáveis pelo cliente

8. GESTÃO DE RISCOS:
   - Identificar riscos técnicos, de segurança e operacionais
   - Pontos de atenção devem incluir ações de mitigação específicas para infraestrutura
   - Considerar riscos de segurança cibernética

9. MATRIZ RACI:
   - Mapear responsabilidades para cada etapa de cada item
   - Incluir equipes técnicas (infraestrutura, segurança, operações)
   - Considerar stakeholders de negócio e técnicos

10. SLAs:
    - Extrair todos os SLAs mencionados no contrato
    - Incluir SLAs de disponibilidade, performance e segurança
    - Incluir penalizações/glosas quando especificadas
    - Definir indicadores mensuráveis (KPIs)

11. COMUNICAÇÃO:
    - Identificar stakeholders técnicos e de negócio
    - Definir frequência e canais apropriados
    - Incluir reuniões técnicas e de gestão

12. FORMATO DE DATAS:
    - Todas as datas no formato YYYY-MM-DD
    - Usar data de assinatura do contrato como início
    - Usar data de término do contrato como fim

13. RETORNO:
    - Retorne APENAS JSON válido, sem texto adicional
    - Garanta que todas as etapas estejam vinculadas aos itens do contrato
    - Cada item deve ter as 4 etapas (Planejamento, Implantação, Execução, Suporte)
    - Durações estimadas devem ser realistas para projetos de infraestrutura
"""
    
    @staticmethod
    def gerar_plano_trabalho_completo(projeto, texto_documento: str) -> Dict[str, Any]:
        """
        Gera um plano de trabalho completo usando IA para um projeto específico
        
        Args:
            projeto: Instância de Projeto (com item_contrato vinculado)
            texto_documento: Texto extraído do documento
            
        Returns:
            Dict com o plano completo
        """
        analyzer = ContractAIAnalyzer()
        contrato = projeto.contrato
        
        # Busca informações do item do contrato vinculado ao projeto
        item_contrato_info = ""
        if projeto.item_contrato:
            item = projeto.item_contrato
            item_contrato_info = f"""
ITEM DO CONTRATO VINCULADO AO PROJETO:
- Número: {item.numero_item}
- Tipo: {item.get_tipo_display()} ({item.tipo})
- Descrição: {item.descricao}
- Quantidade: {item.quantidade} {item.unidade}
- Valor Unitário: R$ {item.valor_unitario}
"""
        
        # Identifica fornecedor principal para modelo de plano
        fornecedor_info = ""
        if contrato.fornecedores:
            fornecedor_principal = contrato.fornecedores[0] if isinstance(contrato.fornecedores, list) else contrato.fornecedores
            fornecedor_info = f"\nFornecedor Principal: {fornecedor_principal}\n"
        
        # Monta contexto otimizado incluindo informações do projeto e item do contrato
        contexto = f"""PROJETO:
- Nome: {projeto.nome}
- Descrição: {projeto.descricao or 'Sem descrição'}
{item_contrato_info}
CONTRATO:
- Número: {contrato.numero_contrato}
- Cliente: {contrato.cliente.nome_razao_social}
- Objeto: {contrato.objeto[:200]}
- Vigência: {contrato.vigencia} meses
- Data de Início: {contrato.data_assinatura}
- Data de Término: {contrato.data_fim_atual or contrato.data_fim}
- Valor: R$ {contrato.valor_inicial}
- Regime Legal: {contrato.get_regime_legal_display()}
{fornecedor_info}
Documento do Contrato:
{texto_documento[:80000]}
"""
        
        # Analisa com IA usando OpenAI
        client = analyzer._get_openai_client()
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": ContractAIService.PLANO_TRABALHO_PROMPT},
                {"role": "user", "content": f"{contexto}\n\nGere o plano de trabalho completo para este contrato."}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        result = response.choices[0].message.content
        
        return json.loads(result.strip())
    
    @staticmethod
    def criar_plano_trabalho(projeto, dados_plano: Dict[str, Any], usuario):
        """
        Cria o PlanoTrabalho a partir dos dados gerados pela IA
        
        Args:
            projeto: Instância de Projeto (já deve ter item_contrato vinculado)
            dados_plano: Dict com dados do plano
            usuario: Usuário que está criando
        """
        from contracts.models import PlanoTrabalho, SLAImportante
        from datetime import datetime
        
        contrato = projeto.contrato
        
        # Identificar fornecedor principal do projeto (baseado no contrato ou item do contrato)
        fornecedor = None
        if contrato.fornecedores:
            fornecedor = contrato.fornecedores[0] if isinstance(contrato.fornecedores, list) and contrato.fornecedores else None
        elif projeto.item_contrato:
            # Tenta identificar fornecedor pelo item do contrato
            # (pode ser necessário buscar em ItemFornecedor relacionado)
            pass
        
        # Define datas padrão: data de assinatura e data de término do contrato
        # Se a IA fornecer datas, tenta usar, senão usa as do contrato
        data_inicio = contrato.data_assinatura or datetime.now().date()
        data_fim = contrato.data_fim_atual or contrato.data_fim or data_inicio
        
        # Se a IA forneceu datas no cronograma, tenta usar (mas valida)
        if 'cronograma' in dados_plano and dados_plano['cronograma'].get('data_inicio_prevista'):
            try:
                data_inicio_ia = datetime.strptime(
                    dados_plano['cronograma']['data_inicio_prevista'], '%Y-%m-%d'
                ).date()
                # Só usa se for válida e não anterior à data de assinatura
                if data_inicio_ia >= contrato.data_assinatura:
                    data_inicio = data_inicio_ia
            except (ValueError, TypeError, KeyError):
                pass  # Mantém o padrão do contrato
        
        if 'cronograma' in dados_plano and dados_plano['cronograma'].get('data_fim_prevista'):
            try:
                data_fim_ia = datetime.strptime(
                    dados_plano['cronograma']['data_fim_prevista'], '%Y-%m-%d'
                ).date()
                # Só usa se for válida e não posterior à data de término do contrato
                if data_fim_ia <= (contrato.data_fim_atual or contrato.data_fim or data_fim):
                    data_fim = data_fim_ia
            except (ValueError, TypeError, KeyError):
                pass  # Mantém o padrão do contrato
        
        # Processa processo_execucao baseado nos itens do contrato
        # Novo formato: etapas por item (não mais sprints/tarefas)
        processo_execucao = dados_plano.get('processo_execucao', [])
        
        # Se não tiver processo_execucao no novo formato, tenta converter do formato antigo
        if not processo_execucao:
            # Tenta converter fluxo_trabalho_fases antigo (compatibilidade)
            fluxo_fases = dados_plano.get('fluxo_trabalho_fases', [])
            if fluxo_fases:
                # Converte para o novo formato por item
                processo_execucao = []
                # Agrupa por item se possível, senão cria estrutura genérica
                for fase_data in fluxo_fases:
                    fase_nome = fase_data.get('fase', '').lower()
                    fase_key = {
                        'planejamento': 'planejamento',
                        'implantação': 'implantacao',
                        'implantacao': 'implantacao',
                        'execução': 'execucao',
                        'execucao': 'execucao',
                        'suporte/sustentação': 'suporte',
                        'suporte': 'suporte',
                        'sustentação': 'suporte'
                    }.get(fase_nome, 'execucao')
                    
                    # Cria etapa genérica (sem vínculo específico a item)
                    processo_execucao.append({
                        'fase': fase_key,
                        'nome': fase_data.get('fase', 'Etapa'),
                        'descricao': fase_data.get('descricao', ''),
                        'objetivo': fase_data.get('objetivo_fase', ''),
                        'duracao_estimada_dias': fase_data.get('duracao_dias', 30),
                        'entregaveis': fase_data.get('entregaveis_principais', []),
                        'atividades_principais': []
                    })
        
        # Se ainda não tiver, usa processo_execucao antigo
        if not processo_execucao:
            processo_execucao = dados_plano.get('processo_execucao', [])
        
        plano = PlanoTrabalho(
            projeto=projeto,
            fornecedor=fornecedor,
            resumo_contrato=dados_plano.get('resumo_contrato', ''),
            pontos_atencao=dados_plano.get('pontos_atencao', []),
            processo_execucao=processo_execucao,
            data_inicio_prevista=data_inicio,
            data_fim_prevista=data_fim,
            cronograma_detalhado=dados_plano.get('cronograma', {}).get('marcos', []),
            plano_comunicacao=dados_plano.get('plano_comunicacao', {}),
            template_status_report=dados_plano.get('template_status_report', ''),
            frequencia_status_report=dados_plano.get('frequencia_status_report', 'semanal'),
            status='pendente_aprovacao',
            criado_por=usuario,
        )
        plano.save()
        
        # Criar Cláusulas Críticas
        from contracts.models import ClausulaCritica
        for clausula_data in dados_plano.get('clausulas_criticas', []):
            prazo = None
            if clausula_data.get('prazo_atencao'):
                try:
                    prazo = datetime.strptime(clausula_data['prazo_atencao'], '%Y-%m-%d').date()
                except:
                    pass
            
            ClausulaCritica.objects.create(
                contrato=contrato,
                titulo=clausula_data.get('titulo', 'Cláusula'),
                numero_clausula=clausula_data.get('numero_clausula'),
                descricao=clausula_data.get('descricao', ''),
                impacto=clausula_data.get('impacto', 'medio'),
                acao_necessaria=clausula_data.get('acao_necessaria', ''),
                prazo_atencao=prazo,
            )
        
        # Criar SLAs importantes com penalizações
        from contracts.models import QuadroPenalizacao
        for sla_data in dados_plano.get('slas_importantes', []):
            sla = SLAImportante.objects.create(
                contrato=contrato,
                nome=sla_data.get('nome', 'SLA'),
                descricao=sla_data.get('descricao', ''),
                tempo_resposta_horas=int(sla_data.get('tempo_resposta_horas', 4)),
                tempo_solucao_horas=int(sla_data.get('tempo_solucao_horas', 24)),
                prioridade=sla_data.get('prioridade', 'media'),
                alerta_antes_horas=int(sla_data.get('alerta_antes_horas', 24)),
            )
            
            # Criar penalizações/glosas
            for penal_data in sla_data.get('penalizacoes', []):
                QuadroPenalizacao.objects.create(
                    sla_importante=sla,
                    descricao=penal_data.get('descricao', ''),
                    tipo=penal_data.get('tipo', 'penalizacao'),
                    percentual=Decimal(str(penal_data.get('percentual', 0))) if penal_data.get('percentual') else None,
                    valor_fixo=Decimal(str(penal_data.get('valor_fixo', 0))) if penal_data.get('valor_fixo') else None,
                    condicao_aplicacao=penal_data.get('condicao_aplicacao', ''),
                )
        
        # Criar Matriz RACI
        from contracts.models import MatrizRACI
        for raci_data in dados_plano.get('matriz_raci', []):
            MatrizRACI.objects.create(
                contrato=contrato,
                atividade=raci_data.get('atividade', ''),
                fase=raci_data.get('fase', 'execucao'),
                responsavel=raci_data.get('responsavel'),
                aprovador=raci_data.get('aprovador'),
                consultado=raci_data.get('consultado'),
                informado=raci_data.get('informado'),
            )
        
        return plano
    
    @staticmethod
    def gerar_sprints_por_item_contrato(contrato, data_inicio_prevista):
        """
        Gera sprints baseadas nos itens do contrato (software, hardware, solução)
        Cada item terá 4 sprints: Planejamento, Implantação, Execução e Suporte
        
        Args:
            contrato: Instância de Contrato
            data_inicio_prevista: Data de início do projeto
            
        Returns:
            Lista de dicionários com dados das sprints
        """
        from datetime import timedelta
        
        sprints_data = []
        itens_produto = contrato.itens.filter(tipo__in=['software', 'hardware', 'solucao'])
        
        logger.info(f"Gerando sprints para {itens_produto.count()} itens do contrato")
        
        if not itens_produto.exists():
            logger.warning("Nenhum item de produto (software/hardware/solução) encontrado no contrato")
            return sprints_data
        
        # Fases do ciclo de vida do item
        FASES = [
            {
                'nome': 'Planejamento',
                'duracao_semanas': 2,
                'tarefas_padrao': [
                    {
                        'titulo': 'Análise de Requisitos',
                        'descricao': 'Análise detalhada dos requisitos do item',
                        'tipo': 'planejamento',
                        'prioridade': 'alta',
                        'horas_planejadas': 16
                    },
                    {
                        'titulo': 'Planejamento Técnico',
                        'descricao': 'Definição da arquitetura e abordagem técnica',
                        'tipo': 'planejamento',
                        'prioridade': 'alta',
                        'horas_planejadas': 24
                    },
                    {
                        'titulo': 'Definição de Escopo',
                        'descricao': 'Delimitação do escopo de trabalho',
                        'tipo': 'planejamento',
                        'prioridade': 'media',
                        'horas_planejadas': 8
                    }
                ]
            },
            {
                'nome': 'Implantação',
                'duracao_semanas': 4,
                'tarefas_padrao': [
                    {
                        'titulo': 'Preparação do Ambiente',
                        'descricao': 'Configuração e preparação do ambiente de implantação',
                        'tipo': 'implantacao',
                        'prioridade': 'alta',
                        'horas_planejadas': 16
                    },
                    {
                        'titulo': 'Instalação/Configuração',
                        'descricao': 'Instalação e configuração do item',
                        'tipo': 'implantacao',
                        'prioridade': 'alta',
                        'horas_planejadas': 32
                    },
                    {
                        'titulo': 'Testes de Implantação',
                        'descricao': 'Testes para validar a implantação',
                        'tipo': 'teste',
                        'prioridade': 'alta',
                        'horas_planejadas': 16
                    },
                    {
                        'titulo': 'Documentação Técnica',
                        'descricao': 'Documentação da implantação realizada',
                        'tipo': 'documentacao',
                        'prioridade': 'media',
                        'horas_planejadas': 8
                    }
                ]
            },
            {
                'nome': 'Execução',
                'duracao_semanas': 8,
                'tarefas_padrao': [
                    {
                        'titulo': 'Ativação em Produção',
                        'descricao': 'Ativação do item em ambiente de produção',
                        'tipo': 'execucao',
                        'prioridade': 'alta',
                        'horas_planejadas': 8
                    },
                    {
                        'titulo': 'Monitoramento Inicial',
                        'descricao': 'Monitoramento e acompanhamento inicial do item em uso',
                        'tipo': 'execucao',
                        'prioridade': 'alta',
                        'horas_planejadas': 16
                    },
                    {
                        'titulo': 'Ajustes e Otimizações',
                        'descricao': 'Ajustes e otimizações baseados no uso real',
                        'tipo': 'execucao',
                        'prioridade': 'media',
                        'horas_planejadas': 24
                    },
                    {
                        'titulo': 'Treinamento de Usuários',
                        'descricao': 'Treinamento dos usuários finais',
                        'tipo': 'treinamento',
                        'prioridade': 'media',
                        'horas_planejadas': 16
                    }
                ]
            },
            {
                'nome': 'Suporte',
                'duracao_semanas': 12,
                'tarefas_padrao': [
                    {
                        'titulo': 'Suporte Contínuo',
                        'descricao': 'Suporte e manutenção contínua do item',
                        'tipo': 'suporte',
                        'prioridade': 'media',
                        'horas_planejadas': 40
                    },
                    {
                        'titulo': 'Monitoramento de Performance',
                        'descricao': 'Monitoramento contínuo da performance e disponibilidade',
                        'tipo': 'suporte',
                        'prioridade': 'media',
                        'horas_planejadas': 24
                    },
                    {
                        'titulo': 'Atualizações e Patches',
                        'descricao': 'Aplicação de atualizações e patches quando necessário',
                        'tipo': 'suporte',
                        'prioridade': 'baixa',
                        'horas_planejadas': 16
                    }
                ]
            }
        ]
        
        dias_acumulados = 0
        
        for item in itens_produto:
            logger.info(f"Processando item: {item.descricao[:50]}... (Tipo: {item.tipo})")
            
            for fase in FASES:
                nome_sprint = f"{item.numero_item} - {item.descricao[:30]}... - {fase['nome']}"
                
                # Calcula datas
                data_inicio_sprint = data_inicio_prevista + timedelta(days=dias_acumulados)
                duracao_dias = fase['duracao_semanas'] * 7
                data_fim_sprint = data_inicio_sprint + timedelta(days=duracao_dias)
                
                # Personaliza tarefas com informações do item
                tarefas = []
                for tarefa_padrao in fase['tarefas_padrao']:
                    tarefa = tarefa_padrao.copy()
                    tarefa['descricao'] = f"{tarefa_padrao['descricao']} - Item: {item.descricao[:50]}"
                    tarefa['entregavel'] = f"{fase['nome']} do item {item.numero_item} - {item.descricao[:50]}"
                    tarefas.append(tarefa)
                
                sprints_data.append({
                    'nome': nome_sprint,
                    'objetivo': f"{fase['nome']} do item {item.numero_item}: {item.descricao[:100]}",
                    'descricao': f"Sprint de {fase['nome']} para o item {item.numero_item} ({item.tipo})",
                    'data_inicio': data_inicio_sprint.isoformat(),
                    'data_fim': data_fim_sprint.isoformat(),
                    'duracao_semanas': fase['duracao_semanas'],
                    'duracao_dias': duracao_dias,
                    'tarefas': tarefas,
                    'item_contrato_id': item.id,
                    'fase': fase['nome'].lower()
                })
                
                # Atualiza dias acumulados (com 1 semana de buffer entre fases do mesmo item)
                dias_acumulados += duracao_dias + 7
            
            # Buffer entre itens (2 semanas)
            dias_acumulados += 14
        
        logger.info(f"Geradas {len(sprints_data)} sprints baseadas nos itens do contrato")
        return sprints_data
    
    @staticmethod
    def criar_projeto_sprints_tarefas(plano, usuario):
        """
        Cria projeto, sprints e tarefas automaticamente após aprovação do plano
        REGRA IMPORTANTE: Cada item de software, hardware ou solução deve ter sprints para:
        - Planejamento
        - Implantação
        - Execução
        - Suporte
        
        O objetivo é que todos os itens cheguem à fase de Suporte (implantados e em uso)
        
        Args:
            plano: Instância de PlanoTrabalho
            usuario: Usuário que aprovou
        """
        from contracts.models import Projeto, Sprint, Tarefa
        from datetime import datetime, timedelta
        from decimal import Decimal
        
        logger.info(f"Iniciando criação de projeto/sprints/tarefas para plano {plano.pk}")
        
        # Busca sprints do plano aprovado
        sprints_data = []
        
        # NOVA LÓGICA: Gera sprints baseadas nos itens do contrato
        itens_produto = plano.contrato.itens.filter(tipo__in=['software', 'hardware', 'solucao'])
        
        if itens_produto.exists():
            logger.info(f"Gerando sprints baseadas em {itens_produto.count()} itens do contrato")
            sprints_data = ContractAIService.gerar_sprints_por_item_contrato(
                plano.contrato,
                plano.data_inicio_prevista
            )
        
        # Fallback: Verifica se processo_execucao existe e não está vazio
        if not sprints_data:
            processo_execucao = plano.processo_execucao
            if not processo_execucao:
                processo_execucao = []
            
            logger.info(f"Processo execução tem {len(processo_execucao)} etapas (usando fallback)")
            
            # Prioriza fluxo_trabalho_fases (formato melhorado com melhores práticas)
            if processo_execucao and len(processo_execucao) > 0:
                # Processa processo_execucao que já contém as sprints estruturadas
                dias_acumulados = 0
                for idx, etapa in enumerate(processo_execucao):
                    logger.info(f"Processando etapa {idx + 1}: {etapa.get('nome', etapa.get('etapa', 'Sem nome'))}")
                # Garante que 'tarefas' é uma lista, mesmo que vazia
                tarefas_etapa = etapa.get('tarefas', []) or []
                
                # Converte as datas para objetos date
                data_inicio_str = etapa.get('data_inicio')
                data_fim_str = etapa.get('data_fim')
                
                data_inicio_sprint = None
                data_fim_sprint = None
                
                if data_inicio_str:
                    try:
                        if isinstance(data_inicio_str, str):
                            data_inicio_sprint = datetime.strptime(data_inicio_str, '%Y-%m-%d').date()
                        else:
                            data_inicio_sprint = data_inicio_str
                    except (ValueError, TypeError):
                        logger.warning(f"Data de início inválida para sprint: {data_inicio_str}")
                
                if data_fim_str:
                    try:
                        if isinstance(data_fim_str, str):
                            data_fim_sprint = datetime.strptime(data_fim_str, '%Y-%m-%d').date()
                        else:
                            data_fim_sprint = data_fim_str
                    except (ValueError, TypeError):
                        logger.warning(f"Data de fim inválida para sprint: {data_fim_str}")

                # Se as datas não puderem ser parseadas, calcula baseado na duração
                if not data_inicio_sprint and plano.data_inicio_prevista:
                    data_inicio_sprint = plano.data_inicio_prevista + timedelta(days=dias_acumulados)
                
                if not data_fim_sprint and data_inicio_sprint:
                    duracao = etapa.get('duracao_dias', 14)  # Padrão de 2 semanas
                    # Se tiver duracao_semanas, converte para dias
                    if 'duracao_semanas' in etapa:
                        duracao = etapa.get('duracao_semanas', 2) * 7
                    data_fim_sprint = data_inicio_sprint + timedelta(days=duracao)
                
                if data_inicio_sprint and data_fim_sprint:
                    sprints_data.append({
                        'nome': etapa.get('nome', etapa.get('etapa', 'Sprint')),
                        'objetivo': etapa.get('objetivo', etapa.get('descricao', '')),
                        'data_inicio': data_inicio_sprint.isoformat(),
                        'data_fim': data_fim_sprint.isoformat(),
                        'tarefas': tarefas_etapa
                    })
                    # Atualiza dias acumulados para próxima sprint
                    dias_acumulados = (data_fim_sprint - plano.data_inicio_prevista).days
                else:
                    logger.warning(f"Não foi possível determinar datas para a sprint: {etapa.get('nome')}")
            
            # Fallback: busca da análise original se processo_execucao estiver vazio
            if not sprints_data:
                analise = plano.contrato.analises_origem.first()
                if analise and analise.dados_extraidos:
                    fluxo_fases = analise.dados_extraidos.get('fluxo_trabalho_fases', [])
                    if fluxo_fases:
                        dias_acumulados = 0
                        for fase_data in fluxo_fases:
                            sprints_fase = fase_data.get('sprints', [])
                            
                            for sprint in sprints_fase:
                                # Calcula datas se não estiverem definidas
                                if not sprint.get('data_inicio') or not sprint.get('data_fim'):
                                    data_inicio = plano.data_inicio_prevista + timedelta(days=dias_acumulados)
                                # Usa duracao_semanas se disponível, senão duracao_dias, senão padrão 14 dias
                                duracao_semanas = sprint.get('duracao_semanas', 2)
                                duracao = sprint.get('duracao_dias', duracao_semanas * 7)
                                data_fim = data_inicio + timedelta(days=duracao)
                                sprint['data_inicio'] = data_inicio.isoformat()
                                sprint['data_fim'] = data_fim.isoformat()
                                dias_acumulados += duracao
                            else:
                                # Se já tem datas, calcula dias acumulados para próxima sprint
                                try:
                                    data_inicio = datetime.strptime(sprint['data_inicio'], '%Y-%m-%d').date()
                                    data_fim = datetime.strptime(sprint['data_fim'], '%Y-%m-%d').date()
                                    dias_acumulados = (data_fim - plano.data_inicio_prevista).days
                                except (ValueError, TypeError) as e:
                                    logger.warning(f"Erro ao parsear datas da sprint: {e}")
                                    data_inicio = plano.data_inicio_prevista + timedelta(days=dias_acumulados)
                                    data_fim = data_inicio + timedelta(days=14)
                                    sprint['data_inicio'] = data_inicio.isoformat()
                                    sprint['data_fim'] = data_fim.isoformat()
                                    dias_acumulados += 14
                            
                            sprints_data.append({
                                'nome': sprint.get('nome', 'Sprint'),
                                'objetivo': sprint.get('objetivo', ''),
                                'data_inicio': sprint.get('data_inicio'),
                                'data_fim': sprint.get('data_fim'),
                                'tarefas': sprint.get('tarefas', [])
                            })
                            dias_acumulados += 1  # 1 dia de buffer entre sprints
            
            # Se ainda não tiver dados, cria sprints básicas
            if not sprints_data:
                logger.warning("Nenhuma sprint encontrada, criando sprint básica")
            # Cria uma sprint básica
            sprints_data.append({
                'nome': 'Sprint Inicial',
                'objetivo': 'Início do projeto',
                'descricao': 'Início do projeto',
                'data_inicio': plano.data_inicio_prevista.isoformat(),
                'data_fim': (plano.data_inicio_prevista + timedelta(days=14)).isoformat(),
                'tarefas': [{
                    'titulo': 'Kick-off do Projeto',
                    'descricao': 'Reunião inicial de alinhamento',
                    'tipo': 'planejamento',
                    'prioridade': 'alta',
                    'horas_planejadas': 8
                }]
            })
        
        logger.info(f"Total de sprints a criar: {len(sprints_data)}")
        
        # Verifica se já existe projeto vinculado
        if plano.projeto:
            logger.info(f"Projeto já existe: {plano.projeto.pk}")
            projeto = plano.projeto
            # Se o projeto já existe, verifica se já tem sprints
            sprints_existentes = projeto.sprints.count()
            if sprints_existentes > 0:
                logger.warning(f"Projeto já tem {sprints_existentes} sprint(s). Pulando criação de novas sprints.")
                return projeto
        else:
            # Criar Projeto
            logger.info("Criando novo projeto")
            projeto = Projeto(
                nome=f"Projeto - {plano.contrato.numero_contrato}",
                descricao=plano.resumo_contrato[:500] if plano.resumo_contrato else '',
                contrato=plano.contrato,
                data_inicio=plano.data_inicio_prevista,
                data_fim_prevista=plano.data_fim_prevista,
                gerente_projeto=plano.contrato.gerente_contrato,
            )
            projeto.save()
            logger.info(f"Projeto criado: {projeto.pk}")
            
            plano.projeto = projeto
            plano.save()
        
        # Criar Sprints e Tarefas
        sprints_criadas = 0
        tarefas_criadas = 0
        
        for idx, sprint_data in enumerate(sprints_data):
            logger.info(f"Criando sprint {idx + 1}/{len(sprints_data)}: {sprint_data.get('nome')}")
            try:
                data_inicio_sprint = datetime.strptime(sprint_data['data_inicio'], '%Y-%m-%d').date()
                data_fim_sprint = datetime.strptime(sprint_data['data_fim'], '%Y-%m-%d').date()
            except (ValueError, KeyError) as e:
                # Se não conseguir parsear, usa datas padrão
                logger.warning(f"Erro ao parsear datas da sprint {sprint_data.get('nome')}: {e}")
                data_inicio_sprint = plano.data_inicio_prevista
                data_fim_sprint = plano.data_inicio_prevista + timedelta(days=14)
            
            try:
                # Define data_fim como a data de término do contrato por padrão
                # Será ajustada futuramente pelo gerente do projeto
                data_fim_contrato = plano.contrato.data_fim_atual or plano.data_fim_prevista
                
                sprint = Sprint(
                    projeto=projeto,
                    nome=sprint_data.get('nome', 'Sprint'),
                    descricao=sprint_data.get('descricao', sprint_data.get('objetivo', '')),
                    data_inicio=data_inicio_sprint,
                    data_fim=data_fim_contrato,  # Usa data de término do contrato por padrão
                )
                sprint.save()
                logger.info(f"Sprint criada: {sprint.pk} - {sprint.nome}")
                sprints_criadas += 1
            except Exception as e:
                logger.error(f"Erro ao criar sprint {sprint_data.get('nome')}: {e}")
                import traceback
                logger.error(traceback.format_exc())
                continue
            
            # Criar Tarefas com melhorias baseadas em melhores práticas
            tarefas_sprint = sprint_data.get('tarefas', [])
            logger.info(f"Sprint {sprint.nome} tem {len(tarefas_sprint)} tarefas")
            
            for tarefa_idx, tarefa_data in enumerate(tarefas_sprint):
                logger.debug(f"Criando tarefa {tarefa_idx + 1}/{len(tarefas_sprint)}: {tarefa_data.get('titulo')}")
                tipo_map = {
                    'desenvolvimento': 'desenvolvimento',
                    'teste': 'teste',
                    'documentacao': 'documentacao',
                    'treinamento': 'treinamento',
                    'implantacao': 'implantacao',
                    'planejamento': 'planejamento',
                    'execucao': 'execucao',
                    'suporte': 'suporte',
                }
                
                prioridade_map = {
                    'critica': 'critica',
                    'alta': 'alta',
                    'media': 'media',
                    'baixa': 'baixa',
                }
                
                # Monta descrição completa incluindo critérios de aceitação e entregáveis
                descricao_completa = tarefa_data.get('descricao', '')
                
                # Adiciona entregável se disponível
                entregavel = tarefa_data.get('entregavel', '')
                if entregavel:
                    descricao_completa += f"\n\nEntregável: {entregavel}"
                
                # Adiciona critérios de aceitação se disponível
                criterios = tarefa_data.get('criterios_aceitacao', '')
                if criterios:
                    descricao_completa += f"\n\nCritérios de Aceitação: {criterios}"
                
                # Adiciona dependências se disponível
                dependencias = tarefa_data.get('dependencias', [])
                if dependencias:
                    descricao_completa += f"\n\nDependências: {', '.join(dependencias)}"
                
                # Determina responsável (pode vir da tarefa ou usar gerente do contrato)
                responsavel_tarefa = plano.contrato.gerente_contrato
                if tarefa_data.get('responsavel'):
                    # Tenta encontrar colaborador pelo nome/função
                    from contracts.models import Colaborador
                    responsavel_nome = tarefa_data.get('responsavel', '').lower()
                    # Busca por nome ou cargo
                    colaborador = Colaborador.objects.filter(
                        Q(nome_completo__icontains=responsavel_nome) |
                        Q(cargo__icontains=responsavel_nome)
                    ).first()
                    if colaborador:
                        responsavel_tarefa = colaborador
                
                horas_planejadas = Decimal(str(tarefa_data.get('horas_planejadas', 40)))
                # Garante mínimo de 8 horas
                if horas_planejadas < 8:
                    horas_planejadas = Decimal('8.00')
                
                try:
                    tarefa = Tarefa(
                        sprint=sprint,
                        titulo=tarefa_data.get('titulo', 'Tarefa'),
                        descricao=descricao_completa.strip(),
                        tipo=tipo_map.get(tarefa_data.get('tipo', '').lower(), 'desenvolvimento'),
                        prioridade=prioridade_map.get(tarefa_data.get('prioridade', '').lower(), 'media'),
                        horas_planejadas=horas_planejadas,
                        responsavel=responsavel_tarefa,
                    )
                    tarefa.save()
                    tarefas_criadas += 1
                except Exception as e:
                    logger.error(f"Erro ao criar tarefa {tarefa_data.get('titulo')}: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    continue
            
            sprints_criadas += 1
        
        logger.info(f"Projeto {projeto.nome} criado com {sprints_criadas} sprints e {tarefas_criadas} tarefas")
        return projeto
    
    @staticmethod
    def sincronizar_projeto_com_plano(plano, usuario, forcar_sincronizacao=False):
        """
        Sincroniza sprints e tarefas do projeto com as mudanças no plano de trabalho
        Cria novas sprints/tarefas, atualiza existentes e remove as que foram deletadas do plano
        
        Args:
            plano: Instância de PlanoTrabalho
            usuario: Usuário que fez a edição
            forcar_sincronizacao: Se True, sincroniza mesmo se o plano não estiver aprovado
        """
        from contracts.models import Projeto, Sprint, Tarefa
        from datetime import datetime, timedelta
        from decimal import Decimal
        
        # Verifica se tem projeto vinculado
        if not plano.projeto:
            # Se o plano está aprovado, cria o projeto
            if plano.status == 'aprovado':
                logger.warning(f"Plano {plano.pk} aprovado mas sem projeto vinculado. Criando projeto...")
                return ContractAIService.criar_projeto_sprints_tarefas(plano, usuario)
            else:
                logger.info(f"Plano {plano.pk} não tem projeto vinculado e não está aprovado. Sincronização não necessária.")
                return None
        
        # Se não estiver aprovado e não forçar sincronização, apenas loga
        if plano.status != 'aprovado' and not forcar_sincronizacao:
            logger.info(f"Plano {plano.pk} não está aprovado. Use forcar_sincronizacao=True para sincronizar mesmo assim.")
            return None
        
        # Se forçar sincronização, continua mesmo se não estiver aprovado
        if forcar_sincronizacao and plano.status != 'aprovado':
            logger.info(f"Forçando sincronização do plano {plano.pk} mesmo não estando aprovado.")
        
        projeto = plano.projeto
        logger.info(f"Sincronizando projeto {projeto.pk} com plano {plano.pk}")
        
        # Processa processo_execucao do plano
        processo_execucao = plano.processo_execucao or []
        
        if not processo_execucao:
            logger.warning(f"Plano {plano.pk} não tem processo_execucao. Nada para sincronizar.")
            return projeto
        
        # Mapeia sprints existentes por nome (usando nome como identificador)
        sprints_existentes = {sprint.nome: sprint for sprint in projeto.sprints.all()}
        
        # Processa cada etapa do processo_execucao
        sprints_criadas = 0
        sprints_atualizadas = 0
        tarefas_criadas = 0
        tarefas_atualizadas = 0
        
        for etapa in processo_execucao:
            nome_sprint = etapa.get('nome', etapa.get('etapa', 'Sprint'))
            tarefas_etapa = etapa.get('tarefas', []) or []
            
            # Parse de datas
            data_inicio_str = etapa.get('data_inicio')
            data_fim_str = etapa.get('data_fim')
            
            data_inicio_sprint = None
            data_fim_sprint = None
            
            if data_inicio_str:
                try:
                    if isinstance(data_inicio_str, str):
                        data_inicio_sprint = datetime.strptime(data_inicio_str, '%Y-%m-%d').date()
                    else:
                        data_inicio_sprint = data_inicio_str
                except (ValueError, TypeError):
                    logger.warning(f"Data de início inválida para sprint: {data_inicio_str}")
            
            if data_fim_str:
                try:
                    if isinstance(data_fim_str, str):
                        data_fim_sprint = datetime.strptime(data_fim_str, '%Y-%m-%d').date()
                    else:
                        data_fim_sprint = data_fim_str
                except (ValueError, TypeError):
                    logger.warning(f"Data de fim inválida para sprint: {data_fim_str}")
            
            # Se não tiver datas, calcula baseado na duração
            if not data_inicio_sprint:
                data_inicio_sprint = plano.data_inicio_prevista
            
            if not data_fim_sprint:
                duracao = etapa.get('duracao_dias', 14)
                if 'duracao_semanas' in etapa:
                    duracao = etapa.get('duracao_semanas', 2) * 7
                data_fim_sprint = data_inicio_sprint + timedelta(days=duracao)
            
            # Verifica se sprint já existe
            if nome_sprint in sprints_existentes:
                # Atualiza sprint existente
                sprint = sprints_existentes[nome_sprint]
                # Define data_fim como a data de término do contrato por padrão
                data_fim_contrato = plano.contrato.data_fim_atual or plano.data_fim_prevista
                
                sprint.descricao = etapa.get('objetivo', etapa.get('descricao', sprint.descricao))
                sprint.data_inicio = data_inicio_sprint
                sprint.data_fim = data_fim_contrato  # Usa data de término do contrato por padrão
                sprint.save()
                sprints_atualizadas += 1
                logger.info(f"Sprint atualizada: {sprint.nome}")
            else:
                # Define data_fim como a data de término do contrato por padrão
                data_fim_contrato = plano.contrato.data_fim_atual or plano.data_fim_prevista
                
                # Cria nova sprint
                sprint = Sprint(
                    projeto=projeto,
                    nome=nome_sprint,
                    descricao=etapa.get('objetivo', etapa.get('descricao', '')),
                    data_inicio=data_inicio_sprint,
                    data_fim=data_fim_contrato,  # Usa data de término do contrato por padrão
                )
                sprint.save()
                sprints_criadas += 1
                logger.info(f"Sprint criada: {sprint.nome}")
            
            # Sincroniza tarefas da sprint
            tarefas_existentes = {tarefa.titulo: tarefa for tarefa in sprint.tarefas.all()}
            
            for tarefa_data in tarefas_etapa:
                titulo_tarefa = tarefa_data.get('titulo', 'Tarefa')
                
                # Mapeamentos
                tipo_map = {
                    'desenvolvimento': 'desenvolvimento',
                    'teste': 'teste',
                    'documentacao': 'documentacao',
                    'treinamento': 'treinamento',
                    'implantacao': 'implantacao',
                    'planejamento': 'planejamento',
                    'execucao': 'execucao',
                    'suporte': 'suporte',
                }
                
                prioridade_map = {
                    'critica': 'critica',
                    'alta': 'alta',
                    'media': 'media',
                    'baixa': 'baixa',
                }
                
                # Monta descrição completa
                descricao_completa = tarefa_data.get('descricao', '')
                
                entregavel = tarefa_data.get('entregavel', '')
                if entregavel:
                    descricao_completa += f"\n\nEntregável: {entregavel}"
                
                criterios = tarefa_data.get('criterios_aceitacao', '')
                if criterios:
                    descricao_completa += f"\n\nCritérios de Aceitação: {criterios}"
                
                dependencias = tarefa_data.get('dependencias', [])
                if dependencias:
                    descricao_completa += f"\n\nDependências: {', '.join(dependencias)}"
                
                # Determina responsável
                responsavel_tarefa = plano.contrato.gerente_contrato
                if tarefa_data.get('responsavel'):
                    from contracts.models import Colaborador
                    responsavel_nome = tarefa_data.get('responsavel', '').lower()
                    colaborador = Colaborador.objects.filter(
                        Q(nome_completo__icontains=responsavel_nome) |
                        Q(cargo__icontains=responsavel_nome)
                    ).first()
                    if colaborador:
                        responsavel_tarefa = colaborador
                
                horas_planejadas = Decimal(str(tarefa_data.get('horas_planejadas', 40)))
                if horas_planejadas < 8:
                    horas_planejadas = Decimal('8.00')
                
                # Verifica se tarefa já existe
                if titulo_tarefa in tarefas_existentes:
                    # Atualiza tarefa existente
                    tarefa = tarefas_existentes[titulo_tarefa]
                    tarefa.descricao = descricao_completa.strip()
                    tarefa.tipo = tipo_map.get(tarefa_data.get('tipo', '').lower(), tarefa.tipo)
                    tarefa.prioridade = prioridade_map.get(tarefa_data.get('prioridade', '').lower(), tarefa.prioridade)
                    tarefa.horas_planejadas = horas_planejadas
                    tarefa.responsavel = responsavel_tarefa
                    tarefa.save()
                    tarefas_atualizadas += 1
                    logger.info(f"Tarefa atualizada: {tarefa.titulo}")
                else:
                    # Cria nova tarefa
                    try:
                        bilhetar_na_os = tarefa_data.get('bilhetar_na_os', True)
                        
                        tarefa = Tarefa(
                            projeto=projeto,  # Projeto obrigatório
                            sprint=sprint,  # Sprint pode ser None (fica no backlog)
                            titulo=titulo_tarefa,
                            descricao=descricao_completa.strip(),
                            prioridade=prioridade_map.get(tarefa_data.get('prioridade', '').lower(), 'media'),
                            horas_planejadas=horas_planejadas,
                            responsavel=responsavel_tarefa,
                            bilhetar_na_os=bilhetar_na_os,
                            status='pendente',
                        )
                        tarefa.save()
                        tarefas_criadas += 1
                        logger.info(f"Tarefa criada: {tarefa.titulo}")
                    except Exception as e:
                        logger.error(f"Erro ao criar tarefa {titulo_tarefa}: {e}")
                        continue
            
            # Remove tarefas que não estão mais no plano
            tarefas_planos_titulos = {t.get('titulo', 'Tarefa') for t in tarefas_etapa}
            tarefas_para_remover = [
                t for t in sprint.tarefas.all() 
                if t.titulo not in tarefas_planos_titulos and t.status not in ['finalizada', 'concluida']
            ]
            for tarefa in tarefas_para_remover:
                logger.info(f"Removendo tarefa {tarefa.titulo} que não está mais no plano")
                tarefa.delete()
        
        # Remove sprints que não estão mais no plano
        sprints_planos_nomes = {e.get('nome', e.get('etapa', 'Sprint')) for e in processo_execucao}
        sprints_para_remover = [
            s for s in projeto.sprints.all() 
            if s.nome not in sprints_planos_nomes and s.status not in ['finalizada', 'faturada']
        ]
        for sprint in sprints_para_remover:
            logger.info(f"Removendo sprint {sprint.nome} que não está mais no plano")
            sprint.delete()
        
        logger.info(
            f"Sincronização concluída: {sprints_criadas} sprints criadas, "
            f"{sprints_atualizadas} atualizadas, {tarefas_criadas} tarefas criadas, "
            f"{tarefas_atualizadas} tarefas atualizadas"
        )
        
        return projeto

